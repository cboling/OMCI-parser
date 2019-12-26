#!/usr/bin/env python
#
# Copyright (c) 2019 - present.  Boling Consulting Solutions (bcsw.net)
#
# Licensed under the Apache License, Version 2.0 (the "License");
# you may not use this file except in compliance with the License.
# You may obtain a copy of the License at
# http://www.apache.org/licenses/LICENSE-2.0
# Unless required by applicable law or agreed to in writing, software
# distributed under the License is distributed on an "AS IS" BASIS,
# WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
# See the License for the specific language governing permissions and
# limitations under the License.
#
from __future__ import (
    absolute_import, division, print_function, unicode_literals
)
import argparse
from docx import Document
import os
import time
import shutil
from parsed_json import ParsedJson
from versions import VersionHeading
from actions import Actions


def parse_args():
    parser = argparse.ArgumentParser(description='G.988 CSV Parser')

    parser.add_argument('--ITU', '-I', action='store',
                        default='T-REC-G.988-201711-I!!MSW-E.docx',
                        help='Path to ITU G.988 specification document')

    parser.add_argument('--input', '-i', action='store',
                        default='G.988.Parsed.json',
                        help='Input parsed data filename, default: G.988.Parsed.json')

    parser.add_argument('--dir', '-d', action='store',
                        default='csv',
                        help='Directory name to place csv-generated files, default: csv')

    parser.add_argument('--force', '-f', action='store_true',
                        help='Overwrite output directory, default is to fail if it exists')

    args = parser.parse_args()
    return args


class Main(object):
    """ Main program """
    def __init__(self):
        self.args = parse_args()
        self.paragraphs = None
        self.parsed = ParsedJson()
        self.version = VersionHeading()
        self.version.name = 'csv-generator'
        self.version.create_time = time.time()
        self.version.itu_document = self.args.ITU

    @staticmethod
    def get_file_hash(filename):
        import hashlib
        with open(filename, 'rb') as f:
            data = f.read()
            return hashlib.sha256(data).hexdigest()

    def load_itu_document(self):
        return Document(self.args.ITU)

    def start(self):
        print("Loading ITU Document '{}' and parsed data file '{}'".format(self.args.ITU,
                                                                           self.args.input))
        try:
            document = self.load_itu_document()
            self.paragraphs = document.paragraphs

            # Directory exists
            if os.path.isdir(self.args.dir):
                if not self.args.force:
                    print("Directory '{}' exists, use --force to overwrite", self.args.dir)
                    return
                shutil.rmtree(self.args.dir)

            # Load JSON class ID list and sort by ID
            self.parsed.load(self.args.input)
            self.parsed.add(self.version)

            class_ids = [c for c in self.parsed.class_ids.values()]
            class_ids.sort(key=lambda x: x.cid)

            # Create output directory
            os.mkdir(self.args.dir)

            # Create the file
            filename = os.path.join(self.args.dir, "omci.csv")
            with open(filename, 'w') as f:
                header = 'Class ID, Class Name, CID Supported, Created By, Actions, MIB Reset, '\
                         'Attribute Number, Attribute Name, Access, Optional, Attr Supported, Type, '\
                         'AVC, Default Value, Constraints, Notes\n'
                f.write(header)

                for class_id in class_ids:
                    name = class_id.name
                    cid = class_id.cid
                    cid_supported = 'n/a'
                    created_by = 'OLT' if Actions.Create in class_id.actions else 'ONU'
                    # actions = '","'.join(action.name for action in class_id.actions)
                    actions = ' '.join(action.name for action in class_id.actions)
                    reset = 'yes' if created_by == 'ONU' else 'no'
                    cid_info = '{}, {}, {}, {}, {}, {}'.format(name, cid, cid_supported,
                                                               created_by, actions, reset)
                    f.write('{}, , , , , , , , , ,\n'.format(cid_info))
                    for attribute in class_id.attributes:
                        pass
                        attr_number = attribute.index
                        attr_name = attribute.name
                        # attr_access = '","'.join(access.name for access in attribute.access)
                        attr_access = ' '.join(access.name for access in attribute.access)
                        attr_optional = 'yes' if attribute.optional else 'no'
                        attr_supported = 'n/a'
                        attr_type = self.attr_type(attribute)
                        attr_default = 'None'
                        attr_constraint = 'None'
                        attr_avc = 'yes' if attribute.avc else 'no'
                        #
                        f.write('{}, {}, {}, {}, {}, {}, {}, {}, {}, {}\n'.format(cid_info, attr_number, attr_name,
                                                                                  attr_access, attr_optional,
                                                                                  attr_supported, attr_type,
                                                                                  attr_avc,
                                                                                  attr_default, attr_constraint))
        except Exception as _e:
            raise
        # Done
        print('CSV generation completed successfully')

    def attr_type(self, attribute):
        results = ''
        pass
        if attribute.table_support:
            results = 'table({})'.format(attribute.size.octets)
            pass
            pass
            pass

        elif attribute.counter:
            results = 'uint'
            pass
        else:
            size = '{}'.format(attribute.size.octets)
            if attribute.size.repeat_count > 1:
                if attribute.size.repeat_max:
                    size += '*{}..{}'.format(attribute.size.repeat_count,
                                             attribute.size.repeat_count)
                else:
                    size += '*{}'.format(attribute.size.repeat_count)

            results = 'uint({})'.format(size)

        return results


if __name__ == '__main__':
    Main().start()
