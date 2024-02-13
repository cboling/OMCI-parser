#!/usr/bin/env python
#
# Copyright (c) 2018 - present.  Boling Consulting Solutions (bcsw.net)
#
# Licensed under the Apache License, Version 2.0 (the "License");
# you may not use this file except in compliance with the License.
# You may obtain a copy of the License at
# http://www.apache.org/licenses/LICENSE-2.0
# Unless required by applicable law or agreed to in writing, software
# distributed under the License is distributed on an "AS IS" BASIS,
# WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
# See the License for the specific language governing permissions and
# limitations under the License.

#
#   This program will walk the ITU G.988 docx document and extract section
#   headers (and following paragraphs) and any Table Items.  It then
#   will save this to a pre-compiled JSON file that can then be used
#   later to quickly load the sections for further processing

from __future__ import (
    absolute_import, division, print_function, unicode_literals
)
import sys
import argparse
import time
from docx import Document
from docx.document import Document as _Document
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.table import _Cell
from docx.table import Table as DocxTable
from docx.text.paragraph import Paragraph

from preparsed_json import PreParsedJson
from parser_lib.section import SectionHeading
from parser_lib.versions import VersionHeading
from parser_lib.tables import Table

#
#  This application takes the G.988 Word document and preparses the sections. The parsing is divided
#  up into several steps as the pre-parsing takes several minutes and is usually on ran the first
#  time or after a bug-fix or feature addition to the preParse.py (this) file.
#
#  The process from scratch is:
#
#   1. Pre-parse the G.988 Word Document via 'preParse.py' to create the 'G.988.Precompiled.json' file
#
#   2. Parse the G.988 JSON via 'parser.py' to create the G.988.Parsed.json file.  At this point,
#      there is just a minimal fragment 'G.988.augment.yaml' file that really as a little bit of data
#
#   3. Run the 'augmentGenerator.py' file to create an 'augmented.yaml' file in the 'metadata' sub-
#      directory. This will have all the newly parsed JSON converted to YAML form.
#
#   4. Hand edit the augmented.yaml file by hand and adjust the values as needed.  Once you are done,
#      overwrite the base directory's 'G.988.augment.yaml' file with this.  You now have a
#      metadata file that will be used everytime you either run the parser (or code generator) with
#      you updated metadata.
#
#   5. Run the 'parser.py' program again. This will pick up the metadata hint you created in step 4
#      and will create an updated 'G.988.Parsed.json' file with your data.
#
#   6. Run either the C or Go code generator to generate your classes.
#
############################################################################
#
#   If the pre-parser is upgraded or a bug is fixed and needs to be ran again due to updates, run steps 1-6
#
#   If the parser is upgraded or a bug is fixed, you only need to run steps 5 & 6
#
#   If the code generator you are interested in is upgraded or fixed, you only need to run step 6
#
############################################################################

def parse_args():
    parser = argparse.ArgumentParser(description='G.988 Pre-process Parser')

    parser.add_argument('--input', '-i', action='store',
                        default='T-REC-G.988-202003-I!Amd3!MSW-E.docx',
                        help='Path to ITU G.988 specification document')

    parser.add_argument('--output', '-o', action='store',
                        default='G.988.PreCompiled.json',
                        help='Output filename, default: G.988.PreCompiled.json')

    parser.add_argument('--verbose', '-v', action='count',
                        default=0,
                        help='Increase verbosity level')

    args = parser.parse_args()
    return args.input, args.output


class Main:
    """ Main program """
    def __init__(self):
        self.args = parse_args()
        self.preparsed = PreParsedJson()

        version = VersionHeading()
        version.name = 'pre-parser'
        version.create_time = time.time()
        version.itu_document = self.args[0]
        version.version = self.get_version()
        version.sha256 = self.get_file_hash(version.itu_document)
        self.preparsed.add(version)

    def verbose(self, level: int = 1) -> bool:
        return level <= self.args.verbose

    @staticmethod
    def get_version():
        with open('VERSION', 'r') as f:
            for line in f:
                line = line.strip().lower()
                if len(line) > 0:
                    return line

    @staticmethod
    def get_file_hash(filename):
        import hashlib
        with open(filename, 'rb') as f:
            data = f.read()
            return hashlib.sha256(data).hexdigest()

    def start(self, source_file, output):
        # sections = SectionList()
        document = Document(source_file)

        paragraphs = document.paragraphs
        doc_sections = document.sections
        styles = document.styles
        tables = document.tables

        print('Number of sections  : {}'.format(len(doc_sections)))
        print('Number of paragraphs: {}'.format(len(paragraphs)))
        print('Number of styles    : {}, {} are built-in styles'.format(len(styles),
                                                                        len([x for x in styles
                                                                             if x.builtin])))
        print('Number of tables    : {}'.format(len(tables)))
        print('Parsing paragraphs & tables to extract high level information.')
        print('This may take a little while (~3-5 minutes)')

        pnum = 0
        tnum = 0
        current_section = None

        def is_section_header(p):
            return (isinstance(p, Paragraph)
                    and len(p.text)
                    and p.style.builtin
                    and 'heading ' in p.style.name.lower())

        for block in Main.iter_block_items(document):
            if isinstance(block, Paragraph):
                if is_section_header(block):
                    # Save of previous
                    current_section = SectionHeading.create(pnum, block)
                    if current_section:
                        self.preparsed.add(current_section)

                elif len(block.text) > 0 and current_section is not None:
                    current_section.add_contents(pnum)

                pnum += 1

            elif isinstance(block, DocxTable):
                if current_section is not None:
                    table = Table.create(tnum, block)
                    current_section.add_contents(table)
                tnum += 1

            else:
                print('Unsupported block type: {}'.format(type(block)))

            if pnum % 25 == 24:
                print('.', end='')
                sys.stdout.flush()
            if pnum % 2000 == 1999:
                print('')

        # Save to file
        print('')
        print('Saving Section parsing information to {}'.format(output))
        self.preparsed.save(output)

        print('Section pre-parsing are complete')
        self.preparsed.dump()

        # Restore and verify
        self.preparsed.load(output)

        print('Dumping data loaded from saved file for verification')
        for section in self.preparsed.sections:
            print('  Section: {} -> {}'.format(section, section.section_points))

    @staticmethod
    def iter_block_items(parent):
        """
        Generate a reference to each paragraph and table child within *parent*,
        in document order. Each returned value is an instance of either Table or
        Paragraph. *parent* would most commonly be a reference to a main
        Document object, but also works for a _Cell object, which itself can
        contain paragraphs and tables.
        """
        if isinstance(parent, _Document):
            parent_elm = parent.element.body
        elif isinstance(parent, _Cell):
            parent_elm = parent._tc
        else:
            raise ValueError("something's not right")

        for child in parent_elm.iterchildren():
            if isinstance(child, CT_P):
                yield Paragraph(child, parent)
            elif isinstance(child, CT_Tbl):
                yield DocxTable(child, parent)


if __name__ == '__main__':
    try:
        inputfile, output = parse_args()
        Main().start(inputfile, output)

    except Exception as _e:
        raise
