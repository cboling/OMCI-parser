#!/usr/bin/env python
#
# Copyright (c) 2018 - present.  Boling Consulting Solutions (bcsw.net)
#
# Licensed under the Apache License, Version 2.0 (the "License");
# you may not use this file except in compliance with the License.
# You may obtain a copy of the License at
# http://www.apache.org/licenses/LICENSE-2.0
# Unless required by applicable law or agreed to in writing, software
# distributed under the License is distributed on an "AS IS" BASIS,
# WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
# See the License for the specific language governing permissions and
# limitations under the License.
#
from __future__ import (
    absolute_import, division, print_function, unicode_literals
)
import argparse
from docx import Document
import jinja2
import os
import time
import shutil
from golang.classIdmap import create_class_id_map
from golang.managedentity import create_managed_entity_file
from golang.basetemplates import create_base_templates
from golang.versionfile import create_version_file
from parser_lib.parsed_json import ParsedJson
from parser_lib.versions import VersionHeading
import base64


#
#  This application takes a G.988 parsed JSON file and generates 'Go' code from templates
#
#  The process from scratch is:
#
#   1. Pre-parse the G.988 Word Document via 'preParse.py' to create the 'G.988.Precompiled.json' file
#
#   2. Parse the G.988 JSON via 'parser.py' to create the G.988.Parsed.json file.  At this point,
#      there is just a minimal fragment 'G.988.augment.yaml' file that really as a little bit of data
#
#   3. Run the 'augmentGenerator.py' file to create an 'augmented.yaml' file in the 'metadata' sub-
#      directory. This will have all the newly parsed JSON converted to YAML form.
#
#   4. Hand edit the augmented.yaml file by hand and adjust the values as needed.  Once you are done,
#      overwrite the base directory's 'G.988.augment.yaml' file with this.  You now have a
#      metadata file that will be used everytime you either run the parser (or code generator) with
#      you updated metadata.
#
#   5. Run the 'parser.py' program again. This will pick up the metadata hint you created in step 4
#      and will create an updated 'G.988.Parsed.json' file with your data.
#
#   6. Run either the C or Go code generator to generate your classes.
#
############################################################################
#
#   If the pre-parser is upgraded or a bug is fixed and needs to be ran again due to updates, run steps 1-6
#
#   If the parser is upgraded or a bug is fixed, you only need to run steps 5 & 6
#
#   If the code generator you are interested in is upgraded or fixed, you only need to run step 6
#
############################################################################
def parse_args():
    parser = argparse.ArgumentParser(description='G.988 Final Parser')

    parser.add_argument('--ITU', '-I', action='store',
                        default='T-REC-G.988-202003-I!Amd3!MSW-E.docx',
                        help='Path to ITU G.988 specification document')

    parser.add_argument('--input', '-i', action='store',
                        default='G.988.Parsed.json',
                        help='Input parsed data filename, default: G.988.Parsed.json')

    parser.add_argument('--dir', '-d', action='store',
                        default='generated',
                        help='Directory name to place code-generated files, default: generated')

    parser.add_argument('--templates', '-t', action='store',
                        default='golang',
                        help='Location of code-generation Jinja2 templates, default: golang')

    parser.add_argument('--force', '-f', action='store_true',
                        help='Overwrite output directory, default is to fail if it exists')

    args = parser.parse_args()
    return args


def zero_b64_string(size):
    return 'toOctets("{}")'.format(str(base64.b64encode(bytearray(size)))[2:-1])


def attribute_bitmask(index):
    if index <= 0:
        return '0x0000'
    return '{0:#0{1}x}'.format(1 << (16 - index), 6)


class Main(object):
    """ Main program """
    def __init__(self):
        self.args = parse_args()
        self.paragraphs = None
        loader = jinja2.FileSystemLoader(searchpath=self.args.templates)
        self.templateEnv = jinja2.Environment(loader=loader,
                                              extensions=['jinja2.ext.loopcontrols'])
        self.templateEnv.filters['zero_b64_string'] = zero_b64_string
        self.templateEnv.filters['attribute_bitmask'] = attribute_bitmask

        self.parsed = ParsedJson()
        self.version = VersionHeading()
        self.version.name = 'code-generator'
        self.version.create_time = time.time()
        self.version.itu_document = self.args.ITU
        self.version.version = self.get_version()
        self.version.sha256 = self.get_file_hash(self.version.itu_document)


    @staticmethod
    def get_version():
        with open('VERSION', 'r') as f:
            for line in f:
                line = line.strip().lower()
                if len(line) > 0:
                    return line

    @staticmethod
    def get_file_hash(filename):
        import hashlib
        with open(filename, 'rb') as f:
            data = f.read()
            return hashlib.sha256(data).hexdigest()

    def load_itu_document(self):
        return Document(self.args.ITU)

    def start(self):
        print("Loading ITU Document '{}' and parsed data file '{}'".format(self.args.ITU,
                                                                           self.args.input))
        try:
            document = self.load_itu_document()
            self.paragraphs = document.paragraphs

            # Directory exists
            if os.path.isdir(self.args.dir):
                if not self.args.force:
                    print("Directory '{}' exists, use --force to overwrite", self.args.dir)
                    return
                shutil.rmtree(self.args.dir)

            # Load JSON class ID list and sort by ID
            self.parsed.load(self.args.input)
            self.parsed.add(self.version)

            class_ids = [c for c in self.parsed.class_ids.values()]
            class_ids.sort(key=lambda x: x.cid)

            # Create output directory
            os.mkdir(self.args.dir)

            # Generate some somewhat fixed templates
            create_base_templates(self.args.dir, self.templateEnv)

            # Create Version File
            create_version_file(self.parsed.versions, self.args.dir, self.templateEnv)

            # Create Class ID Map
            create_class_id_map(class_ids, self.args.dir, self.templateEnv)

            # Create Managed Entity files
            for class_id in class_ids:
                create_managed_entity_file(class_id, self.args.dir, self.templateEnv, self.paragraphs)
                # if class_id.cid == 2:
                #    print('TODO: remove later')  # Good place for breakpoint

        except Exception as _e:
            raise
        # Done
        print('Code generation completed successfully')


if __name__ == '__main__':
    Main().start()
