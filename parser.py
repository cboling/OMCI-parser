#!/usr/bin/env python
#
# Copyright (c) 2018 - present.  Boling Consulting Solutions (bcsw.net)
#
# Licensed under the Apache License, Version 2.0 (the "License");
# you may not use this file except in compliance with the License.
# You may obtain a copy of the License at
# http://www.apache.org/licenses/LICENSE-2.0
# Unless required by applicable law or agreed to in writing, software
# distributed under the License is distributed on an "AS IS" BASIS,
# WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
# See the License for the specific language governing permissions and
# limitations under the License.
#
from __future__ import (
    absolute_import, division, print_function, unicode_literals
)

import argparse
import copy
import time

from docx import Document

from parser_lib.attributes import Attribute
from parser_lib.class_id import ClassIdList, ClassAccess
from parser_lib.parsed_json import ParsedJson
from parser_lib.parsed_yaml import MetadataYAML
from parser_lib.text import camelcase
from parser_lib.versions import VersionHeading
from preparsed_json import PreParsedJson

#
#  This application takes the pre-parsed JSON file and further parses it to output suitable for
#  running a code generator with.
#
#  The process from scratch is:
#
#   1. Pre-parse the G.988 Word Document via 'preParse.py' to create the 'G.988.Precompiled.json' file
#
#   2. Parse the G.988 JSON via 'parser.py' to create the G.988.Parsed.json file.  At this point,
#      there is just a minimal fragment 'G.988.augment.yaml' file that really as a little bit of data
#
#   3. Run the 'augmentGenerator.py' file to create an 'augmented.yaml' file in the 'metadata' sub-
#      directory. This will have all the newly parsed JSON converted to YAML form.
#
#   4. Hand edit the augmented.yaml file by hand and adjust the values as needed.  Once you are done,
#      overwrite the base directory's 'G.988.augment.yaml' file with this.  You now have a
#      metadata file that will be used every time you either run the parser (or code generator) with
#      your updated metadata.
#
#   5. Run the 'parser.py' program again. This will pick up the metadata hint you created in step 4
#      and will create an updated 'G.988.Parsed.json' file with your data.
#
#   6. Run either the C or Go code generator to generate your classes.
#
############################################################################
#
#   If the pre-parser is upgraded or a bug is fixed and needs to be ran again due to updates, run steps 1-6
#
#   If the parser is upgraded or a bug is fixed, you only need to run steps 5 & 6
#
#   If the code generator you are interested in is upgraded or fixed, you only need to run step 6
#
############################################################################

MEClassSection = "11.2.4"       # Class IDs


def parse_args():
    parser = argparse.ArgumentParser(description='G.988 Parser')

    parser.add_argument('--ITU', '-I', action='store',
                        default='T-REC-G.988-202003-I!Amd3!MSW-E.docx',
                        help='Path to ITU G.988 specification document')

    parser.add_argument('--input', '-i', action='store',
                        default='G.988.PreCompiled.json',
                        help='Path to pre-parsed G.988 data, default: G.988.PreCompiled.json')

    parser.add_argument('--hints', '-H', action='store',
                        default='G.988.augment.yaml',
                        help='Path to hand modified G.988 input data, default: G.988.augment.yaml')

    parser.add_argument('--output', '-o', action='store',
                        default='G.988.Parsed.json',
                        help='Output filename, default: G.988.Parsed.json')

    parser.add_argument('--classes', '-c', action='store',
                        default='11.2.4',
                        help='Document section number with ME Class IDs, default: 11.2.4')

    args = parser.parse_args()
    return args


class Main(object):
    """ Main program """
    def __init__(self):
        self.args = parse_args()
        self.paragraphs = None
        self.body = None

        self.preparsed = PreParsedJson()
        self.parsed = ParsedJson()
        self.attribute_hints = dict()       # Class ID -> attribute hints list
        version = VersionHeading()
        version.name = 'parser'
        version.create_time = time.time()
        version.itu_document = self.args.ITU
        version.version = self.get_version()
        version.sha256 = self.get_file_hash(version.itu_document)
        self.parsed.add(version)

    @staticmethod
    def get_version():
        with open('VERSION', 'r') as f:
            for line in f:
                line = line.strip().lower()
                if len(line) > 0:
                    return line

    @staticmethod
    def get_file_hash(filename):
        import hashlib
        with open(filename, 'rb') as f:
            data = f.read()
            return hashlib.sha256(data).hexdigest()

    @property
    def sections(self):
        return self.preparsed.sections

    @property
    def section_list(self):
        return self.preparsed.section_list

    @property
    def class_ids(self):
        return self.parsed.class_ids

    def load_itu_document(self):
        return Document(self.args.ITU)

    def start(self):
        print("Loading ITU Document '{}' and parsed data file '{}'".format(self.args.ITU,
                                                                           self.args.input))
        self.preparsed.load(self.args.input)
        for version in self.preparsed.versions:
            self.parsed.add(version)

        self.attribute_hints = MetadataYAML().load(self.args.hints)

        document = self.load_itu_document()
        self.paragraphs = document.paragraphs
        # doc_sections = document.sections
        # styles = document.styles
        # self.body = document.element.body

        print('Extracting ME Class ID values')
        class_ids = ClassIdList.parse_sections(self.section_list,
                                               self.args.classes)
        for _cid, me_class in class_ids.items():
            self.parsed.add(me_class)

        print('Found {} ME Class ID entries. {} have sections associated to them'.
              format(len(self.class_ids),
                     len([c for c in self.class_ids.values()
                          if c.section is not None])))

        crazy_formatted_mes = {}  # Try all
        todo_class_ids = {k: v for k, v in self.class_ids.items()
                          if k not in crazy_formatted_mes}

        print('Managed Entities without Sections')
        for c in [c for c in todo_class_ids.values() if c.section is None]:
            print('    {:>4}: {}'.format(c.cid, c.name))

        # Work with what we have
        todo_class_ids = {cid: c for cid, c in todo_class_ids.items()
                          if c.section is not None}
        print('')
        print('Parsing deeper for managed Entities with Sections')

        final_class_ids = todo_class_ids

        # c = final_class_ids[441]              # Uncomment for fast debug of a single Class ID
        # c.deep_parse(self.paragraphs)

        for c in final_class_ids.values():
            if c.section is None:
                c.failure(None, None)
                continue

            print('    {:>9}:  {:>4}: {} -> {}'.format(c.section.section_number,
                                                       c.cid,
                                                       c.name,
                                                       camelcase(c.name)))
            c.deep_parse(self.paragraphs)
            pass

        # Some just need some manual intervention
        final_class_ids = self.fix_difficult_class_ids(final_class_ids)

        # Who creates them
        final_class_ids = self.find_class_access(final_class_ids)

        completed = len([c for c in final_class_ids.values() if c.state == 'complete'])
        failed = len([c for c in final_class_ids.values() if c.state == 'failure'])

        print('Of {} MEs, {} were parsed successfully and {} failed'.format(len(final_class_ids),
                                                                            completed,
                                                                            failed))
        # Run some sanity checks
        print('\n\n\nValidating ME Class Information, total of {}:\n'.
              format(len(final_class_ids)))

        class_with_issues = dict()
        class_with_no_actions = dict()
        class_with_no_attributes = dict()
        attributes_with_no_access = dict()
        attributes_with_no_size = dict()
        attributes_with_zero_size = dict()
        class_with_too_many_attributes = dict()
        num_attributes = 0

        for c in final_class_ids.values():
            print('  ID: {}: {} -\t{}'.format(c.cid, c.section.section_number, c.name),
                  end='')

            if c.state != 'complete':
                print('\t\tParsing ended in state {}', c.state)
                class_with_issues[c.cid] = c

            if len(c.actions) == 0:
                print('\t\tActions: No actions decoded for ME')
                class_with_issues[c.cid] = c
                class_with_no_actions[c.cid] = c
                c.failure(None, None)       # Mark invalid
            else:
                print('\t\tActions: {}'.format({a.name for a in c.actions}))

            if len(c.attributes) == 0:
                print('\t\tNO ATTRIBUTES')      # TODO Look for 'set' without 'get'
                class_with_issues[c.cid] = c
                class_with_no_attributes[c.cid] = c
                c.failure(None, None)       # Mark invalid

            elif len(c.attributes) > 17:        # Entity ID counts as well in this list
                print('\t\tTOO MANY ATTRIBUTES')
                class_with_issues[c.cid] = c
                class_with_too_many_attributes[c.cid] = c
                c.failure(None, None)       # Mark invalid

            else:
                for attr in c.attributes:
                    num_attributes += 1
                    print('\t\t\t\t{}'.format(attr.name), end='')
                    if attr.access is None or len(attr.access) == 0:
                        print('\t\t\t\tNO ACCESS INFORMATION')
                        attributes_with_no_access[c.cid] = c
                        c.failure(None, None)       # Mark invalid
                    else:
                        print('\t\t\t\tAccess: {}'.format({a.name for a in attr.access}))

                    if attr.size is None:
                        attributes_with_no_size[c.cid] = c
                        print('\t\t\t\tNO SIZE INFORMATION')
                        c.failure(None, None)       # Mark invalid

                    elif attr.size.octets == 0:
                        attributes_with_zero_size[c.cid] = c
                        print('\t\t\t\tSIZE zero')
                        c.failure(None, None)       # Mark invalid

        print('Section parsing is complete, saving JSON output...')
        print('=======================================================')

        # Output the results to a JSON file so it can be used by a code-generation
        # tool
        self.parsed.save(self.args.output)

        # Restore and verify
        self.parsed.load(self.args.output)
        self.parsed.dump()

        # Results
        print("Of the {} class IDs, {} had issues: {} had no actions and {} had no attributes and {} with too many".
              format(len(final_class_ids), len(class_with_issues), len(class_with_no_actions),
                     len(class_with_no_attributes), len(class_with_too_many_attributes)))

        print("Of the {} attributes, {} had no access info and {} had no size info and {} had zero size".
              format(num_attributes, len(attributes_with_no_access), len(attributes_with_no_size),
                     len(attributes_with_zero_size)))

        bad_cids = {key for key in class_with_no_actions.keys()} | \
                   {key for key in class_with_no_attributes.keys()} | \
                   {key for key in class_with_too_many_attributes.keys()} | \
                   {key for key in attributes_with_no_access.keys()} | \
                   {key for key in attributes_with_no_size.keys()} | \
                   {key for key in attributes_with_zero_size.keys()}

        if bad_cids:
            print('=======================================================')
            print("Bad Classes: {}".format(len(bad_cids)))
            for cid in bad_cids:
                c = final_class_ids[cid]
                bad = []
                if cid in class_with_no_actions:
                    bad.append("No Action")
                if cid in class_with_no_attributes:
                    bad.append("No Attributes")
                if cid in class_with_too_many_attributes:
                    bad.append("Too Nany Attributes")
                if cid in attributes_with_no_access:
                    bad.append("No Access")
                if cid in attributes_with_no_size:
                    bad.append("No Size")
                if cid in attributes_with_zero_size:
                    bad.append("Zero Size")

                print("  {:3d}: {:<60s}: {}".format(cid, c.name, ', '.join(what for what in bad)))

    def fix_difficult_class_ids(self, class_list):
        # Special exception. Ethernet frame performance monitoring history data downstream
        # is in identical upstream and only a note of that exists. Fix it now
        from parser_lib.actions import Actions
        from parser_lib.size import AttributeSize
        from parser_lib.attributes import AttributeAccess, AttributeList, AttributeType

        # Circuit Pack is now created only by ONU (OLT creation kept only for backwards
        # compatibility
        if 6 in class_list.keys():
            item = class_list[6]
            item.access = ClassAccess.CreatedByOnu

        if 58 in class_list.keys():
            item = class_list[58]
            item.actions.add(Actions.GetNext)

        if 113 in class_list.keys():
            item = class_list[113]
            item.alarms._alarms[7] = ('leftr defect seconds', item.alarms._alarms[7][1])

        if 134 in class_list.keys():
            item = class_list[134]
            item.actions.add(Actions.Test)

        if 408 in class_list.keys():
            item = class_list[408]
            item.alarms._alarms[0] = ('leftr defect seconds', item.alarms._alarms[0][1])

        if 149 in class_list.keys():
            sip = class_list[149]       # SIP config portal (zero size for config text table)

            sz = AttributeSize()
            sz._octets = 25             # Assume it is 25 at most since it is vendor specific
            sip.attributes[1].size = sz

        if 150 in class_list.keys():
            item = class_list[150]
            item.actions.add(Actions.GetNext)

        if 154 in class_list.keys():
            mgc = class_list[154]       # MGC config portal (zero size for config text table)

            sz = AttributeSize()
            sz._octets = 25             # Assume it is 25 at most since it is vendor specific
            mgc.attributes[1].size = sz

        # For SIP user data, the Username&Password attribute is a pointer
        # to a security methods ME and is 2 bytes but is in the document as
        # just (2)
        if 153 in class_list.keys():
            sip = class_list[153]

            sz = AttributeSize()
            sz._octets = 25
            sip.attributes[3].size = sz
            sip.attributes[3].access.add(AttributeAccess.Read)
            sip.attributes[3].access.add(AttributeAccess.Write)

            sz2 = AttributeSize()
            sz2._octets = 2
            sip.attributes[4].size = sz2

        # Large string
        if 157 in class_list.keys():
            item = class_list[157]
            part1 = item.attributes[2]
            part1.name = 'Part 1'
            part1.optional = False
            part1.size._repeat_max = 1

            for part in range(2, 16):
                newPart = copy.deepcopy(part1)
                newPart.name = 'Part {}'.format(part)
                item.attributes.add(newPart)

        # ONU remote debug - reply table size not bounded
        if 158 in class_list.keys():
            item = class_list[158]
            reply_table = item.attributes[3]
            sz = copy.deepcopy(reply_table.size)
            sz._octets = -1
            reply_table.size = sz

        # MCAST GEM Interworking - IPv4
        if 281 in class_list.keys():
            item = class_list[281]
            sz = AttributeSize()
            sz._octets = 12
            sz.getnext_required = True
            item.attributes[9].size = sz
            item.attributes[9].access.add(AttributeAccess.Read)
            item.attributes[9].access.add(AttributeAccess.Write)

        # OMCI.  IPv6 Table is 24N octets
        if 287 in class_list.keys():
            item = class_list[287]
            item.attributes[1].getnext_required = True

            sz = AttributeSize()
            sz._octets = 1
            sz.getnext_required = True
            item.attributes[2].size = sz

        # Managed entity tables.  4 tables need fixing
        if 288 in class_list.keys():
            me = class_list[288]
            class_list[288].name += ' ME'       # To avoid conflicts with Go file/struct names
            sz = AttributeSize()
            sz._octets = 1
            me.attributes[4].size = sz
            me.attributes[5].size = sz

        # Managed entity code points table.  Table is 2*n octets
        if 289 in class_list.keys():
            class_list[289].name += ' ME'       # To avoid conflicts with Go file/struct names

        # Dot1ag maintenance domain. Has multiple attributes defined on one line that need
        # to be split up
        if 299 in class_list.keys():
            item = class_list[299]
            for index in range(6, 3, -1):
                item.attributes[index] = item.attributes[index-1]

            item.attributes[3].name = "MD Name 1"
            item.attributes[4].name = "MD Name 2"

        # Dot1ag maintenance association. Has multiple attributes defined on one line that need
        # to be split up
        if 300 in class_list.keys():
            item = class_list[300]
            for index in range(7, 3, -1):
                item.attributes[index] = item.attributes[index-1]

            item.attributes[3].name = "Short MA Name 1"
            item.attributes[4].name = "Short MA Name 2"

        # Dot1ag chassis-managment info. Has multiple attributes defined on one line that need
        # to be split up.  Three times...
        if 306 in class_list.keys():
            item = class_list[306]
            for index in range(8, 3, -1):
                item.attributes[index] = item.attributes[index-1]

            item.attributes[3].name = "Chassis ID Part 1"
            item.attributes[4].name = "Chassis ID Part 2"

            for index in range(9, 6, -1):
                item.attributes[index] = item.attributes[index-1]

            item.attributes[6].name = "Management Address Domain 1"
            item.attributes[7].name = "Management Address Domain 2"

            for index in range(10, 9, -1):
                item.attributes[index] = item.attributes[index-1]

            item.attributes[9].name = "Management Address 1"
            item.attributes[10].name = "Management Address 2"

        # Octet String - 1..15 row entries
        if 307 in class_list.keys():
            item = class_list[307]
            part1 = item.attributes[2]
            part1.name = 'Part 1'
            part1.optional = False
            part1.size._repeat_max = 1

            for part in range(2, 16):
                newPart = copy.deepcopy(part1)
                newPart.name = 'Part {}'.format(part)
                newPart.optional = True
                item.attributes.add(newPart)

        # General Purpose Buffer - Buffer table is one big unbounded string
        if 308 in class_list.keys():
            item = class_list[308]
            buffer_table = item.attributes[2]
            sz = copy.deepcopy(buffer_table.size)
            sz._octets = -1
            buffer_table.size = sz

        # For multicast operations profile - Attributes getting polluted with table info/descriptions
        if 309 in class_list.keys():
            item = class_list[309]
            if len(item.attributes) == 20:
                first_bad = 8   # Table Control is description of attribute 7
                next_good = 11  # Pick back up to good attributes at the Static Access Control List Table
                # Also attribute 7 lost its information on access...
                old_attributes = item.attributes
                item.attributes = AttributeList()
                for index, attribute in enumerate(old_attributes):
                    if index < first_bad or index >= next_good:
                        item.attributes.add(attribute)

                # And a quick attribute name fixup here
                item.attributes[16].name = "Downstream IGMP and multicast TCI"
                # Some enumerated items
                item.attributes[1].attribute_type = AttributeType.Enumeration
                item.attributes[2].attribute_type = AttributeType.Enumeration
                item.attributes[3].attribute_type = AttributeType.Enumeration

                item.attributes[16].name = "Downstream IGMP and multicast TCI"
                # The pain in the rear table in the document
                sz = AttributeSize()
                sz._octets = 24
                sz.getnext_required = True
                item.attributes[7].size = sz
                item.attributes[7].access.add(AttributeAccess.Read)
                item.attributes[7].access.add(AttributeAccess.Write)

        # For multicast subscriber config info. very hard to decode automatically
        if 310 in class_list.keys():
            msci = class_list[310]
            msci.attributes.remove(8)
            item = msci.attributes[7]
            sz = AttributeSize()
            sz._octets = 22
            sz.getnext_required = True
            item.size = sz
            item.access.add(AttributeAccess.Read)
            item.access.add(AttributeAccess.Write)

        if 321 in class_list.keys() and 322 in class_list.keys():
            down = class_list[321]
            up = class_list[322]
            down.attributes = up.attributes
            down.actions = up.actions
            down.optional_actions = up.optional_actions
            down.alarms = up.alarms
            down.avcs = up.avcs
            down.test_results = up.test_results
            down.hidden = up.hidden

        # xDSL line inventory and status data part 5
        if 325 in class_list.keys():
            item = class_list[325]
            try:
                # Type in document, not table attributes present
                item.actions.remove(Actions.GetNext)
            except KeyError:
                pass

        # Enhanced security control
        if 332 in class_list.keys():
            item = class_list[332]
            try:
                # Type in document, not table attributes present
                item.attributes[7].access.add(AttributeAccess.Read)
            except KeyError:
                pass

        # ONU dynamic power management control
        if 336 in class_list.keys():
            item = class_list[336]
            try:
                old_attributes = item.attributes
                item.attributes = AttributeList()
                for index, attribute in enumerate(old_attributes):
                    if index < 11:
                        item.attributes.add(attribute)

                item.attributes.add(old_attributes[12])
                item.attributes.add(old_attributes[14])
                sz = AttributeSize()
                sz._octets = 1
                item.attributes[10].size = sz
                item.attributes[10].optional = True
                item.attributes[10].access.add(AttributeAccess.Read)

                item.attributes[11].size = sz
                item.attributes[11].optional = True
                item.attributes[11].access.add(AttributeAccess.Read)
                item.attributes[11].access.add(AttributeAccess.Write)
            except KeyError:
                pass

        # xDSL line inventory and status data part 8
        if 414 in class_list.keys():
            item = class_list[414]
            try:
                # Type in document, not table attributes present
                item.actions.remove(Actions.GetNext)
            except KeyError:
                pass

        # Fast Channel COnfigureation Profile ME is missing Managed Entity
        if 432 in class_list.keys():
            item = class_list[432]
            for index in range(11, 0, -1):
                item.attributes[index] = item.attributes[index - 1]

            item.attributes[0] = Attribute().load(
                {
                    "name":          "Managed Entity Id",
                    "description":   [
                        5374
                    ],
                    "access":        [
                        "Read",
                        "SetByCreate"
                    ],
                    "optional":      False,
                    "deprecated":    False,
                    "size":          {
                        "octets":           2,
                        "bits":             None,
                        "repeat_count":     1,
                        "repeat_max":       1,
                        "getnext_required": False
                    },
                    "avc":           False,
                    "tca":           False,
                    "table-support": False,
                    "type":          "Pointer",
                    "constraint":    None,
                    "default":       None
                }, 0)

        # ONU-3G
        if 441 in class_list.keys():
            item = class_list[441]

            old_attributes = item.attributes
            item.attributes = AttributeList()
            for index, attribute in enumerate(old_attributes):
                if index != 3:
                    item.attributes.add(attribute)

            sz = AttributeSize()
            sz._octets = 1
            item.attributes[2].size = sz
            item.attributes[2].optional = False
            item.attributes[2].access.add(AttributeAccess.Read)

            sz = copy.deepcopy(item.attributes[8].size)
            sz._octets = 25     # N bytes. Vendor specific. 25 bytes is max get attribute size
            item.attributes[8].size = sz
            item.attributes[6].size = sz        # Same issue M rows of 'n' size....

        # Now even some other crazy things
        class_list = self.fix_other_difficulties(class_list)

        # Find counter and other types of attributes
        self.find_attribute_types(class_list)

        return class_list

    @staticmethod
    def fix_other_difficulties(class_list):
        # Some uncommon cleanups
        # for cid, cls in class_list.items():
        #     pass
        return class_list

    @staticmethod
    def find_class_access(class_list):
        from parser_lib.class_id import ClassAccess, Actions
        for _, item in class_list.items():
            if item.access == ClassAccess.UnknownAccess and len(item.actions):
                if Actions.Create in item.actions:
                    item.access = ClassAccess.CreatedByOlt
                else:
                    item.access = ClassAccess.CreatedByOnu

        return class_list

    def find_attribute_types(self, class_list):
        # A bit more in depth look at the attributes
        from parser_lib.attributes import AttributeType

        for cid, item in class_list.items():
            # Any hints to apply before seeding default settings?
            if cid in self.attribute_hints:
                hints = self.attribute_hints[cid]
                for index, attr_hint in hints['attributes'].items():
                    item_attribute = next((attr for attr in item.attributes if attr.index == index), None)
                    if item_attribute is not None:
                        if 'type' in attr_hint:
                            item_attribute.attribute_type = attr_hint['type']
                        if 'default' in attr_hint:
                            item_attribute.default = attr_hint['default']
                        if 'constraint' in attr_hint:
                            item_attribute.constraint = attr_hint['constraint']

            for attr in item.attributes:
                # Only set unknown types (or integer since it may be a counter)
                if attr.attribute_type == AttributeType.Unknown:
                    attr.find_type(item)

        return class_list


if __name__ == '__main__':
    try:
        Main().start()

    except Exception as _e:
        raise
