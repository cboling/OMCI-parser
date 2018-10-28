#!/usr/bin/env python
#
# Copyright (c) 2018 - present.  Boling Consulting Solutions (bcsw.net)
#
# Licensed under the Apache License, Version 2.0 (the "License");
# you may not use this file except in compliance with the License.
# You may obtain a copy of the License at
# http://www.apache.org/licenses/LICENSE-2.0
# Unless required by applicable law or agreed to in writing, software
# distributed under the License is distributed on an "AS IS" BASIS,
# WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
# See the License for the specific language governing permissions and
# limitations under the License.
#
from __future__ import (
    absolute_import, division, print_function, unicode_literals
)
import argparse
from docx import Document
import json
import jinja2
import os
import shutil
from golang.classIdMap import create_class_id_map


def parse_args():
    parser = argparse.ArgumentParser(description='G.988 Final Parser')

    parser.add_argument('--ITU', '-I', action='store',
                        default='ITU-T G.988-201711.docx',
                        help='Path to ITU G.988 specification document')

    parser.add_argument('--input', '-i', action='store',
                        default='G.988.Parsed.json',
                        help='Input parsed data filename, default: G.988.Parsed.json')

    parser.add_argument('--dir', '-d', action='store',
                        default='generated',
                        help='Directory name to place code-generated files, default: generated')

    parser.add_argument('--force', '-f', action='store_true',
                        help='Overwrite output directory, default is to fail if it exists')

    args = parser.parse_args()
    return args


class Main(object):
    """ Main program """
    def __init__(self):
        self.args = parse_args()
        self.paragraphs = None
        self.class_ids = None
        loader = jinja2.FileSystemLoader(searchpath='./golang')
        self.templateEnv = jinja2.Environment(loader=loader)

    def load_itu_document(self):
        return Document(self.args.ITU)

    def start(self):
        print("Loading ITU Document '{}' and parsed data file '{}'".format(self.args.ITU,
                                                                           self.args.input))
        try:
            document = self.load_itu_document()
            self.paragraphs = document.paragraphs

            # Directory exists
            if os.path.isdir(self.args.dir):
                if not self.args.force:
                    print("Directory '{}' exists, use --force to overwrite", self.args.dir)
                    return
                shutil.rmtree(self.args.dir)

            # Load JSON class id
            with open(self.args.input) as f:
                self.class_ids = json.load(f)

            # Create output directory
            os.mkdir(self.args.dir)

            # Create ClassID to Type map
            create_class_id_map(self.class_ids, self.args.dir, self.templateEnv)

        except Exception as _e:
            pass


if __name__ == '__main__':
    Main().start()
